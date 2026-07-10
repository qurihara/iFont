#!/usr/bin/env python3
# iFont 論文草稿を情報処理学会論文誌ふうの2段組(2カラム)docxに組版するスクリプト。
# build_paper.py からの変更点:
#   (1) 論文の構造化ソースを一時パスではなく paper/paper_content.json から読む(再現性)。
#   (2) 図番号を「本文での初出順」に振り直す(投稿規程の通し番号順に合わせる)。
#   (3) 図の画像パスを実在ファイルに修正し、図7(階層モデル縮小図)も確実に貼り込む。
#   (4) タイトルと概要は1段抜き、本文は2段組。図表は段抜き(両段ぶち抜き)で初出位置に配置。
import os, re, json
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
SRC  = os.path.join(HERE, "paper_content.json")
DEST = os.path.join(ROOT, "iFont論文草稿_想定結果版_情報処理学会論文誌.docx")
FIGDIR = os.path.join(HERE, "figs")
SHRINK = os.path.join(ROOT, "階層モデルとは_縮小の図.png")

# ソース側の図id -> 画像ファイル(内容と画像は一体なのでソースのidで対応づけ、下で本文初出順の新番号に付け替える)。
OLD_FILE = {"図1": os.path.join(FIGDIR, "fig_concept.png"),
            "図2": os.path.join(FIGDIR, "fig_matrix.png"),
            "図3": os.path.join(FIGDIR, "fig_psycho.png"),
            "図4": os.path.join(FIGDIR, "fig_mcd.png"),
            "図5": os.path.join(FIGDIR, "fig_recon.png"),
            "図6": os.path.join(FIGDIR, "fig_algos.png"),
            "図7": SHRINK,
            "図8": os.path.join(FIGDIR, "fig_g.png")}

r = json.load(open(SRC, encoding="utf-8"))

# --- 図番号の振り直し(旧番号 -> 新番号)を「本文での初出順」から自動計算する。
# 手で順序を決めると1段落に複数の図参照がある箇所で誤るため、実際の走査順に連番を振る。
_first = []
for _s in r['sections']:
    for _p in _s['paragraphs']:
        for _fid in re.findall(r'【(図[0-9]+)】', _p):
            if _fid not in _first:
                _first.append(_fid)
for _fid in [f['id'] for f in r['figures']]:   # 本文未参照の図があれば末尾に回す
    if _fid not in _first:
        _first.append(_fid)
REMAP = {old: f"図{i+1}" for i, old in enumerate(_first)}
# 新番号 -> 画像ファイル。
FIGFILE = {REMAP[old]: path for old, path in OLD_FILE.items() if old in REMAP}
print("図番号 振り直し(旧->新):", REMAP)

# --- ソース中の図番号(本文参照・図キャプションのid)を新番号に付け替える ---
def remap_fig_tokens(text):
    return re.sub(r'【(図[0-9]+)】', lambda m: "【" + REMAP.get(m.group(1), m.group(1)) + "】", text)

def strip_abstract_refs(text):
    # 概要(情報処理学会論文誌のスタイルでは自己完結が原則)から図表参照を除去する。
    # 図表番号は本文の初出順で振るため、概要に参照が残ると通し番号が崩れて見える。
    text = re.sub(r'（(?:【(?:図|表)[0-9]+】)+）', '', text)  # 参照だけの全角括弧を丸ごと削除
    text = re.sub(r'\((?:【(?:図|表)[0-9]+】)+\)', '', text)  # 半角括弧版
    text = re.sub(r'【(?:図|表)[0-9]+】', '', text)           # 裸で残った参照も削除
    return text

r['abstractJa'] = strip_abstract_refs(r['abstractJa'])
r['abstractEn'] = strip_abstract_refs(r['abstractEn'])
for s in r['sections']:
    s['paragraphs'] = [remap_fig_tokens(p) for p in s['paragraphs']]
for f in r['figures']:
    f['id'] = REMAP.get(f['id'], f['id'])
r['figures'].sort(key=lambda f: int(re.sub(r'\D', '', f['id'])))

figcap = {f['id']: f['caption'] for f in r['figures']}
tabs   = {t['id']: t for t in r['tables']}

# --- 書式ヘルパ ---
JP = "Yu Mincho"; JPG = "Yu Gothic"; ASC = "Times New Roman"
def setrun(rn, size=9, bold=False, jp=JP, asc=ASC, color=None):
    rn.font.size = Pt(size); rn.font.bold = bold; rn.font.name = asc
    if color: rn.font.color.rgb = RGBColor.from_string(color)
    rp = rn._element.get_or_add_rPr(); rf = rp.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rp.append(rf)
    rf.set(qn('w:eastAsia'), jp); rf.set(qn('w:ascii'), asc); rf.set(qn('w:hAnsi'), asc)

doc = Document()
st = doc.styles['Normal']; st.font.name = ASC; st.font.size = Pt(9)
st.element.rPr.rFonts.set(qn('w:eastAsia'), JP)

def apply_page(section, ncols, space_mm=8.0):
    """A4・余白・段数をセクションに適用する。"""
    section.page_width = Mm(210); section.page_height = Mm(297)
    section.top_margin = Mm(20); section.bottom_margin = Mm(20)
    section.left_margin = Mm(18); section.right_margin = Mm(18)
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols'); sectPr.append(cols)
    cols.set(qn('w:num'), str(ncols))
    cols.set(qn('w:space'), str(int(Mm(space_mm).twips)))
    cols.set(qn('w:equalWidth'), "1")

# セクション0(タイトル・概要)は1段。
apply_page(doc.sections[0], 1)
current = {"cols": 1}
def ensure_cols(n):
    """必要なときだけ連続セクション区切りを入れて段数を切り替える。"""
    if current["cols"] == n:
        return
    sec = doc.add_section(WD_SECTION.CONTINUOUS)
    apply_page(sec, n)
    current["cols"] = n

def para(text, size=9, jp=JP, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True,
         after=4, bold=False, asc=ASC, before=0):
    p = doc.add_paragraph(); p.alignment = align; pf = p.paragraph_format
    pf.space_after = Pt(after); pf.space_before = Pt(before); pf.line_spacing = 1.3
    if indent: pf.first_line_indent = Pt(size)
    setrun(p.add_run(text), size=size, bold=bold, jp=jp, asc=asc); return p

def heading(text, lv):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_before = Pt(8 if lv == 1 else 5); pf.space_after = Pt(3)
    setrun(p.add_run(text), size=11 if lv == 1 else 10, bold=True, jp=JPG, asc="Arial")
    return p

def caption(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format; pf.space_after = Pt(6); pf.space_before = Pt(2)
    setrun(p.add_run(text), size=8.5, jp=JP); return p

placed_fig = set(); placed_tab = set()
def add_figure(fid):
    if fid in placed_fig or fid not in FIGFILE or not os.path.exists(FIGFILE[fid]):
        return
    placed_fig.add(fid)
    ensure_cols(1)                      # 図は両段ぶち抜き(1段セクション)で置く
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    try:
        p.add_run().add_picture(FIGFILE[fid], width=Mm(158))
    except Exception:
        setrun(p.add_run(f"[{fid} 画像読込失敗]"))
    caption(f"{fid}  {figcap.get(fid, '')}")

def add_table(tid):
    if tid in placed_tab or tid not in tabs:
        return
    placed_tab.add(tid); T = tabs[tid]; rows = T['rows']
    if not rows:
        return
    ensure_cols(1)                      # 表も段抜きで置く
    caption(f"{tid}  {T['caption']}")
    tb = doc.add_table(rows=len(rows), cols=max(len(x) for x in rows))
    tb.style = 'Table Grid'; tb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            c = tb.cell(i, j); c.paragraphs[0].clear()
            setrun(c.paragraphs[0].add_run(str(cell)), size=8.5, bold=(i == 0), jp=JP)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# --- タイトル・著者・概要(1段抜き) ---
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
setrun(p.add_run(r['titleJa']), size=15, bold=True, jp=JPG, asc="Arial")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
setrun(p.add_run(r['titleEn']), size=11, bold=False, jp=JP, asc=ASC)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
setrun(p.add_run(r.get('authorsJa', '著者名')), size=11, jp=JP)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
setrun(p.add_run(r.get('affilJa', '') + "（所属は投稿時に確定）"), size=9.5, jp=JP)
doc.add_paragraph().paragraph_format.space_after = Pt(2)
heading("概要", 2); para(r['abstractJa'], size=9, indent=True, after=3)
para("キーワード: " + "，".join(r['keywordsJa']), size=8.5, indent=False, after=6, jp=JP)
heading("Abstract", 2); para(r['abstractEn'], size=9, indent=True, after=3, jp=ASC)
para("Keywords: " + ", ".join(r['keywordsEn']), size=8.5, indent=False, after=8, asc=ASC, jp=ASC)

# --- 本文(2段組)。図表は【図N】【表N】の初出位置に段抜きで配置 ---
FIGREF = re.compile(r'【(図[0-9]+)】'); TABREF = re.compile(r'【(表[0-9]+)】')
for s in r['sections']:
    h = s['heading']
    lv = 2 if re.match(r'^[0-9]+\.[0-9]', h) else (1 if re.match(r'^[0-9]+\.', h) else 2)
    ensure_cols(2); heading(h, lv)
    for pt in s['paragraphs']:
        ensure_cols(2)
        para(pt, indent=not pt.strip().startswith('('))
        for fid in FIGREF.findall(pt): add_figure(fid)
        for tid in TABREF.findall(pt): add_table(tid)

# 念のため未配置の図表を末尾に(通常は全て初出位置で配置済み)
ensure_cols(2)
for fid in ["図1", "図2", "図3", "図4", "図5", "図6", "図7", "図8"]: add_figure(fid)
for tid in ["表1", "表2"]: add_table(tid)

# --- 参考文献(2段組) ---
ensure_cols(2); heading("参考文献", 1)
for ref in r['references']:
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.left_indent = Pt(18); pf.first_line_indent = Pt(-18)
    pf.space_after = Pt(2); pf.line_spacing = 1.1
    setrun(p.add_run(ref), size=8.5, jp=JP)

doc.save(DEST)
print("saved:", DEST)
print("sections:", len(r['sections']), "figs placed:", len(placed_fig),
      "tables placed:", len(placed_tab), "doc sections:", len(doc.sections))
