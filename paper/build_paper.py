#!/usr/bin/env python3
import sys, json, re, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTF="/private/tmp/claude-501/-Users-kurihara-Library-CloudStorage-GoogleDrive-qurihara-gmail-com--------share-google-desktop-share/0a1ebcb0-4b09-4e57-b801-3ba115d17474/tasks/wrpg0127b.output"
FIG="/private/tmp/claude-501/-Users-kurihara-Library-CloudStorage-GoogleDrive-qurihara-gmail-com--------share-google-desktop-share/0a1ebcb0-4b09-4e57-b801-3ba115d17474/scratchpad/paper_figs"
SHRINK="/Users/kurihara/Library/CloudStorage/GoogleDrive-qurihara@gmail.com/マイドライブ/share/google_desktop_share/階層モデルとは_縮小の図.png"
NEW="/Users/kurihara/Library/CloudStorage/GoogleDrive-qurihara@gmail.com/マイドライブ/share/google_desktop_share/iFont"
DEST=NEW+"/iFont論文草稿_想定結果版_情報処理学会論文誌.docx"

FIGFILE={"図1":FIG+"/fig_model.png","図2":FIG+"/fig_matrix.png","図3":FIG+"/fig_psycho.png",
         "図4":FIG+"/fig_algos.png","図5":FIG+"/fig_mcd.png","図6":FIG+"/fig_recon.png","図7":SHRINK}

d=json.loads(open(OUTF).read()); r=d.get('result',d)
if isinstance(r,str): r=json.loads(r)

JP="Yu Mincho"; JPG="Yu Gothic"
def setrun(rn,size=10.5,bold=False,jp=JP,asc="Times New Roman",color=None):
    rn.font.size=Pt(size); rn.font.bold=bold; rn.font.name=asc
    if color: rn.font.color.rgb=RGBColor.from_string(color)
    rp=rn._element.get_or_add_rPr(); rf=rp.find(qn('w:rFonts'))
    if rf is None: rf=rp.makeelement(qn('w:rFonts'),{}); rp.append(rf)
    rf.set(qn('w:eastAsia'),jp); rf.set(qn('w:ascii'),asc); rf.set(qn('w:hAnsi'),asc)
doc=Document()
sec=doc.sections[0]
for s in ["top","bottom","left","right"]: setattr(sec,f"{s}_margin",Pt(54))
st=doc.styles['Normal']; st.font.name="Times New Roman"; st.font.size=Pt(10.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'),JP)

def para(text,size=10.5,jp=JP,align=WD_ALIGN_PARAGRAPH.JUSTIFY,indent=True,after=5,bold=False,asc="Times New Roman",before=0):
    p=doc.add_paragraph(); p.alignment=align; pf=p.paragraph_format
    pf.space_after=Pt(after); pf.space_before=Pt(before); pf.line_spacing=1.4
    if indent: pf.first_line_indent=Pt(10.5)
    setrun(p.add_run(text),size=size,bold=bold,jp=jp,asc=asc); return p
def heading(text,lv):
    p=doc.add_paragraph(); pf=p.paragraph_format; pf.space_before=Pt(9 if lv==1 else 6); pf.space_after=Pt(3)
    setrun(p.add_run(text),size=12.5 if lv==1 else 11,bold=True,jp=JPG,asc="Arial"); return p
def caption(text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; pf=p.paragraph_format; pf.space_after=Pt(8); pf.space_before=Pt(2)
    setrun(p.add_run(text),size=9,jp=JP); return p

placed_fig=set(); placed_tab=set()
figcap={f['id']:f['caption'] for f in r.get('figures',[])}
tabs={t['id']:t for t in r.get('tables',[])}

def add_figure(fid):
    if fid in placed_fig or fid not in FIGFILE or not os.path.exists(FIGFILE[fid]): return
    placed_fig.add(fid)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(4)
    try: p.add_run().add_picture(FIGFILE[fid], width=Inches(4.7))
    except Exception as e: setrun(p.add_run(f"[{fid} 画像読込失敗]"))
    caption(f"{fid}  {figcap.get(fid,'')}")
def add_table(tid):
    if tid in placed_tab or tid not in tabs: return
    placed_tab.add(tid); T=tabs[tid]; rows=T['rows']
    if not rows: return
    caption(f"{tid}  {T['caption']}")
    tb=doc.add_table(rows=len(rows),cols=max(len(x) for x in rows)); tb.style='Table Grid'
    tb.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for i,row in enumerate(rows):
        for j,cell in enumerate(row):
            c=tb.cell(i,j); c.paragraphs[0].clear()
            setrun(c.paragraphs[0].add_run(str(cell)),size=9,bold=(i==0),jp=JP)
    doc.add_paragraph().paragraph_format.space_after=Pt(4)

# タイトル
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; setrun(p.add_run(r['titleJa']),size=15,bold=True,jp=JPG,asc="Arial")
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; setrun(p.add_run(r['titleEn']),size=11,bold=False,jp=JP,asc="Times New Roman")
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; setrun(p.add_run("著者名 所属（投稿時に記入）"),size=10,jp=JP)
doc.add_paragraph().paragraph_format.space_after=Pt(2)
# 概要
heading("概要",2); para(r['abstractJa'],size=9.5,indent=True,after=3)
para("キーワード: "+"，".join(r['keywordsJa']),size=9,indent=False,after=6,jp=JP)
heading("Abstract",2); para(r['abstractEn'],size=9.5,indent=True,after=3,jp="Times New Roman")
para("Keywords: "+", ".join(r['keywordsEn']),size=9,indent=False,after=8,asc="Times New Roman",jp="Times New Roman")

FIGREF=re.compile(r'【(図[0-9])】'); TABREF=re.compile(r'【(表[0-9])】')
for s in r['sections']:
    h=s['heading']; lv=1 if re.match(r'^[0-9]+\.\s', h) or re.match(r'^[0-9]+\.[^0-9]', h) else 2
    if re.match(r'^[0-9]+\.[0-9]', h): lv=2
    heading(h,lv)
    for pt in s['paragraphs']:
        # 箇条書き風(先頭が(1)等)は字下げなし
        para(pt, indent=not pt.strip().startswith('('))
        for fid in FIGREF.findall(pt): add_figure(fid)
        for tid in TABREF.findall(pt): add_table(tid)
# 未配置の図表を末尾に
for fid in ["図1","図2","図3","図4","図5","図6","図7"]: add_figure(fid)
for tid in ["表1","表2"]: add_table(tid)
# 参考文献
heading("参考文献",1)
for ref in r['references']:
    p=doc.add_paragraph(); pf=p.paragraph_format; pf.left_indent=Pt(22); pf.first_line_indent=Pt(-22); pf.space_after=Pt(2); pf.line_spacing=1.15
    setrun(p.add_run(ref),size=9,jp=JP)

doc.save(DEST); print("saved:",DEST); print("sections",len(r['sections']),"figs",len(placed_fig),"tabs",len(placed_tab))
