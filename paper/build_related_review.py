#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# iFont の「短時間の文字提示・連続提示の心理物理」に関する関連研究整理を、
# 情報処理学会論文誌ふうの2段組 docx として単独出力する。
# 注意: メインの「iFont論文草稿_想定結果版_情報処理学会論文誌.docx」には一切触れない。独立ファイルを出力する。
import os
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DEST = os.path.join(ROOT, "iFont関連研究_短時間文字提示と連続提示_情報処理学会論文誌.docx")

TITLE = "短時間の文字提示と連続提示の心理物理 ―iFont の提示設計に関する関連研究の整理―"
AUTHORS = "丸山礼華　栗原一貴"
AFFIL = "津田塾大学"

JP = "Yu Mincho"; JPG = "Yu Gothic"; ASC = "Times New Roman"

def setrun(rn, size=9, bold=False, jp=JP, asc=ASC):
    rn.font.size = Pt(size); rn.font.bold = bold; rn.font.name = asc
    rp = rn._element.get_or_add_rPr(); rf = rp.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rp.append(rf)
    rf.set(qn('w:eastAsia'), jp); rf.set(qn('w:ascii'), asc); rf.set(qn('w:hAnsi'), asc)

doc = Document()
st = doc.styles['Normal']; st.font.name = ASC; st.font.size = Pt(9)
st.element.rPr.rFonts.set(qn('w:eastAsia'), JP)

def apply_page(section, ncols, space_mm=8.0):
    section.page_width = Mm(210); section.page_height = Mm(297)
    section.top_margin = Mm(20); section.bottom_margin = Mm(20)
    section.left_margin = Mm(18); section.right_margin = Mm(18)
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols'); sectPr.append(cols)
    cols.set(qn('w:num'), str(ncols)); cols.set(qn('w:space'), str(int(Mm(space_mm).twips))); cols.set(qn('w:equalWidth'), "1")

apply_page(doc.sections[0], 1)
current = {"cols": 1}
def ensure_cols(n):
    if current["cols"] == n: return
    sec = doc.add_section(WD_SECTION.CONTINUOUS); apply_page(sec, n); current["cols"] = n

def para(text, size=9, indent=True, after=4, jp=JP, asc=ASC, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(); p.alignment = align; pf = p.paragraph_format
    pf.space_after = Pt(after); pf.line_spacing = 1.3
    if indent: pf.first_line_indent = Pt(size)
    setrun(p.add_run(text), size=size, jp=jp, asc=asc); return p

def heading(text, lv=1):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_before = Pt(8 if lv == 1 else 5); pf.space_after = Pt(3)
    setrun(p.add_run(text), size=11 if lv == 1 else 10, bold=True, jp=JPG, asc="Arial"); return p

# --- タイトル(1段) ---
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
setrun(p.add_run(TITLE), size=14, bold=True, jp=JPG, asc="Arial")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
setrun(p.add_run(AUTHORS), size=11, jp=JP)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
setrun(p.add_run(AFFIL + "（所属は投稿時に確定）"), size=9.5, jp=JP)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
setrun(p.add_run("本稿は、iFont 論文本体とは独立の作業メモ（関連研究の整理）である。"), size=9, jp=JP)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

SECS = [
("1. 本稿の目的", [
 "iFont は、かな1文字を固定領域に約0.2秒だけ提示し、時間ゲート提示（透明度上昇やぼかしの解除）で徐々に鮮明化して、明瞭度に対する認識率曲線を測る。さらに2文字を連続提示する条件や、将来的に文字を次々に提示する運用を想定する。本稿では、こうした「短時間の文字提示」と「連続提示における前後の文字の相互干渉」に関する心理物理の先行研究を整理し、iFont の提示設計を支持する点、問題提起となる点、それを踏まえた提示速度の設計方針をまとめる。",
]),
("2. 短時間提示による文字・単語の識別", [
 "静止した文字や単語を短時間だけ提示し、提示時間に対する識別成績を測る手法（タキストスコープ提示）は、Sperling の古典[1]以来確立している。Sperling は5〜500msの提示で文字配列を見せ、全報告では約4〜5文字しか報告できないが、部分報告法では直後に9文字以上が一過的に「利用可能」であることを示した[1]。Petersen と Andersen は、露出時間に対する文字識別の心理測定関数を明示的に定式化し、視覚的注意の理論に露出時間依存性を組み込んで単一・全・部分報告課題のデータに適合させている[2]。",
 "識別成績は提示時間とともに向上して飽和する。McAnany は露出時間24ms〜1sで文字視力を測り、視力は約260msまで改善し、以降はほぼ一定になる（利用時間 utilization time ≈ 260ms）ことを報告した[3]。したがって、固定領域に約0.2秒提示して認識率曲線を測るという iFont の操作は、単一文字の識別成績が飽和する手前から近傍を測るものであり、方法論的に妥当である。",
 "日本語文字については、Sasanuma らが仮名語・漢字語を短時間提示して視野（半球）差を含む認識特性を比較し[4]、Shimizu らが仮名・漢字の正像・鏡像を短時間提示で識別する特性を報告している[5]。大西は、言語への近似度を操作した文字系列の認知・記憶を日本語で検討した[6]。ただし、これらは半球差や言語近似度が主眼であり、「単一のかなを何ms提示すれば識別できるか」という提示時間対認識率の曲線を単一かなについて直接報告したものではない。単一かなの提示時間の定量は、当面は文字一般の知見[1][2][3]からの外挿として扱い、iFont 自身が実測して補うべき対象である。",
]),
("3. 一目で把握できる範囲（visual span）", [
 "一度の固視で確実に認識できる文字数には限界がある。Sperling の把握の範囲（span of apprehension）は約4〜5文字である[1]。Legge らは、一瞥で確実に認識できる文字数を visual span と呼び、これが読書速度を律速する感覚的ボトルネックであると提案した。中心視で約10文字、偏心15°で約1.7文字にまで減少する[7][8]。iFont は固定領域に1文字を提示するため visual span の制約は直接には効かないが、複数文字を同時または近接提示する運用に広げる際には、提示位置・偏心・字間（混み合い）を一定に保たないと単一文字の成績自体が変わる点に注意が要る。",
]),
("4. 連続提示における前後の文字の相互干渉", [
 "文字を連続提示すると、前後の文字が互いの認知に干渉しうる。iFont の2文字課題や将来のRSVP的運用は、まさにこの干渉の影響下にある。",
]),
("4.1 RSVPと注意の瞬き", [
 "文字や画像を同一位置に高速で連続提示する手法を RSVP（Rapid Serial Visual Presentation）という。Raymond らは、RSVP で1つ目の標的を報告すると、その後およそ200〜500ms以内に来る2つ目の標的の報告が著しく低下する「注意の瞬き（attentional blink）」を初めて定量した[9]。Dux と Marois の総説によれば、この落ち込みはおおむね提示順で2〜5番目に現れ、直後（lag 1）はむしろ保たれる lag-1 sparing が生じる[10]。Potter は、RSVP で場面は約100msで「理解」されるが、その記憶表象が後続画像の概念的なマスクに耐えるにはさらに約300msの処理を要することを示した[11]。近年の Potter らは13〜80ms/枚でも名前で指定した標的の検出が偶然を上回ることを示したが、これは意味の粗い検出であって、各文字の完全な同定を保証するものではない[12]。",
]),
("4.2 視覚マスキング", [
 "後続の刺激が先行刺激の見えを抑制する現象を（後方）視覚マスキングという。Breitmeyer と Ögmen は、その主要なモデルと知見を比較・整理した[13]。Enns と Di Lollo は、標的に触れない4点や標的消失後も残る囲みが標的の見えを消す「object substitution masking（共通開始マスキング）」を紹介し、注意が標的に速く向けばマスキングはほぼ消えると論じた[14]。これは連続提示・RSVP的状況に直結する知見である。",
]),
("4.3 視覚的残存とアイコニックメモリ", [
 "刺激が消えた後も、その見えや情報は短時間残存する。Coltheart は、残存を神経的残存・見えの残存（visible persistence）・情報的残存（アイコニックメモリ）に区別し、見えの残存は提示時間が長いほど短くなる（逆持続）などの性質を整理した[15]。Averbach と Coriell は、視覚短期貯蔵の保持が約1/4秒（≈250ms）で減衰し、新しい情報の入力が既存情報を消去（erasure＝マスキング）する局所的機構を示した[16]。すなわち、前の文字の残像が約250ms残り、次の文字の入力がそれを上書きしうる。",
]),
("4.4 語優位効果（成績の非加算性）", [
 "文字が意味のあるまとまり（語）を成すと、単独の文字識別からは予測できない上振れが生じる。Reicher と Wheeler は、同じ文字でも語の中に置かれた方が単独より正確に識別される「語優位効果（word superiority effect）」を示した[17][18]。これは、連続提示された複数文字の識別成績が、単一文字成績の単純な足し合わせにならない場合があることを意味する。",
]),
("5. iFont への含意", [
 "以上の知見を iFont の設計に照らすと、支持・問題提起・設計方針の3つが導かれる。",
]),
("5.1 支持される点", [
 "1文字を固定領域に短時間提示して識別率を提示量（時間）の関数として測るという iFont の中核は、Sperling 以来の標準手法[1]であり、心理測定関数として定式化した先行研究[2]とも整合する。約0.2秒という提示は、単一文字識別が飽和する利用時間（≈260ms）[3]の手前から近傍にあたり、劣化・時間ゲートに対する認識率曲線を測る操作として妥当である。",
]),
("5.2 問題提起される点と対策", [
 "一方、iFont が2文字を連続提示し、将来的に文字を次々に提示する運用へ広げるとき、前後の文字の相互干渉が原理的に予想される。後方マスキング[13][14]、注意の瞬き[9][10]、アイコニックな残存の上書き[15][16]、および語優位効果による非加算性[17][18]がそれである。とくに 0.2秒（200ms）という提示間隔は、見えの残存（≈250ms）や注意の瞬きの窓（200〜500ms）とちょうど重なるため、連続提示では文字どうしが干渉しやすい。これは、iFont が「共調音に相当する視覚的な前後干渉」を測ろうとする2文字課題にとってはむしろ狙いどおりだが、後述する「1文字結果の単純合成」を素朴に適用することは難しくする。対策としては、干渉を評価・統制するために提示間隔（SOA/ISI）を独立変数として掃引し、干渉が消える臨界間隔を実測することが有効である。",
]),
("5.3 「独立タイミングなら単純合成でよい」の成立条件", [
 "「前の文字の影響を受けないタイミングで次を出せば、1文字課題の結果の単純な組合せでよい」という主張は、原理的には支持される。マスキングも注意の瞬きも十分に長い間隔で消えるため、間隔を空ければ各文字はほぼ独立に扱える[13][10]。ただし条件がつく。第一に、ごく短い間隔では lag-1 sparing[10]や時間的統合が起き、2文字が独立でなく融合しうる。第二に、単純合成が成り立つのは各文字を同一の空間条件（固定領域・偏心・字間）で提示する場合に限られ、visual span[7][8]が変われば単一文字成績自体が変わる。第三に、文字が語を成すと語優位効果[17][18]で成績が非加算的に上振れし、意味処理・記憶固定化にはオフ後さらに約300msを要する[11]。したがって、「十分に長い間隔・孤立した無意味文字・固定位置」の条件下では単純合成は概ね妥当だが、短い間隔・語・位置変化の下では条件がつき破綻しうる、というのが文献上の到達点である。",
]),
("5.4 提示速度を可変にする場合の妥当なレンジ", [
 "以上より、提示速度を0.2秒に固定するのではなく、可変の独立変数として実験することが妥当である。単一文字の識別に限れば、成績は露出時間とともに上昇し概ね150〜260msで飽和する[2][3]ため、単文字条件の掃引としては閾近傍の約30〜50msから飽和域の約250〜300msまで（例: 50・100・150・200・250msの離散点）が文献的に妥当なレンジである。連続提示・RSVP文脈では、意味の粗検出は13〜80ms/枚でも部分的に可能だが[12]、各文字を確実に同定させるには前後干渉の窓を避ける必要があり、読書RSVPでも概ね100〜300ms/文字が標準帯である。すなわち、単文字識別率カーブ用に約30〜300ms、連続提示の干渉評価用に約100〜700msのSOAという二段の可変レンジを設定するのが、文献的に無理のない設計である。",
]),
("6. まとめと採用方針", [
 "短時間の文字提示で識別率曲線を測るという iFont の中核は、心理物理の標準手法に支持される[1][2][3]。一方、2文字以上の連続提示では前後干渉が原理的に生じるため、0.2秒固定という前提は再検討を要する。本稿の整理から、次の方針を採用する。(1) 提示速度（1文字あたりの時間、および文字間のSOA）を固定値でなく独立変数として扱い、単文字用に約30〜300ms、連続提示の干渉評価用に約100〜700msのSOAを掃引する。(2) 干渉が十分小さくなる臨界SOAを実測し、そのときに限って「1文字課題の結果の単純合成」で系列の識別を近似できるかを検証する。(3) 2文字課題は、干渉そのもの（視覚版の共調音に相当）を測る条件として引き続き位置づける。これにより、単純合成の成立範囲と、干渉を明示的に扱うべき範囲とを、実データで切り分ける。",
]),
]

REFS = [
 "[1] Sperling, G.: The information available in brief visual presentations, Psychological Monographs: General and Applied, Vol.74, No.11, pp.1–29 (1960).",
 "[2] Petersen, A. and Andersen, T. S.: The effect of exposure duration on visual character identification in single, whole, and partial report, Journal of Experimental Psychology: Human Perception and Performance, Vol.38, No.2, pp.498–514 (2012).",
 "[3] McAnany, J. J.: The effect of exposure duration on visual acuity for letter optotypes and gratings, Vision Research, Vol.105, pp.86–91 (2014).",
 "[4] Sasanuma, S., Itoh, M., Mori, K. and Kobayashi, Y.: Tachistoscopic recognition of kana and kanji words, Neuropsychologia, Vol.15, No.4–5, pp.547–553 (1977).",
 "[5] Shimizu, A., Endo, M. and Nakamura, I.: Tachistoscopic recognition of normal and mirror images of Kana and Kanji characters, Folia Psychiatrica et Neurologica Japonica, Vol.37, No.1, pp.77–84 (1983).",
 "[6] 大西誠一郎: 日本語への近似度を異にする文字系列の認知と記憶について I, 心理学研究, Vol.30, No.5, pp.309–316 (1960).",
 "[7] Legge, G. E., Mansfield, J. S. and Chung, S. T. L.: Psychophysics of reading. XX. Linking letter recognition to reading speed in central and peripheral vision, Vision Research, Vol.41, No.6, pp.725–743 (2001).",
 "[8] Legge, G. E., Cheung, S.-H., Yu, D., Chung, S. T. L., Lee, H.-W. and Owens, D. P.: The case for the visual span as a sensory bottleneck in reading, Journal of Vision, Vol.7, No.2, Article 9 (2007).",
 "[9] Raymond, J. E., Shapiro, K. L. and Arnell, K. M.: Temporary suppression of visual processing in an RSVP task: An attentional blink?, Journal of Experimental Psychology: Human Perception and Performance, Vol.18, No.3, pp.849–860 (1992).",
 "[10] Dux, P. E. and Marois, R.: The attentional blink: A review of data and theory, Attention, Perception, & Psychophysics, Vol.71, No.8, pp.1683–1700 (2009).",
 "[11] Potter, M. C.: Short-term conceptual memory for pictures, Journal of Experimental Psychology: Human Learning and Memory, Vol.2, No.5, pp.509–522 (1976).",
 "[12] Potter, M. C., Wyble, B., Hagmann, C. E. and McCourt, E. S.: Detecting meaning in RSVP at 13 ms per picture, Attention, Perception, & Psychophysics, Vol.76, No.2, pp.270–279 (2014).",
 "[13] Breitmeyer, B. G. and Ögmen, H.: Recent models and findings in visual backward masking: A comparison, review, and update, Perception & Psychophysics, Vol.62, No.8, pp.1572–1595 (2000).",
 "[14] Enns, J. T. and Di Lollo, V.: What's new in visual masking?, Trends in Cognitive Sciences, Vol.4, No.9, pp.345–352 (2000).",
 "[15] Coltheart, M.: Iconic memory and visible persistence, Perception & Psychophysics, Vol.27, No.3, pp.183–228 (1980).",
 "[16] Averbach, E. and Coriell, A. S.: Short-term memory in vision, Bell System Technical Journal, Vol.40, No.1, pp.309–328 (1961).",
 "[17] Reicher, G. M.: Perceptual recognition as a function of meaningfulness of stimulus material, Journal of Experimental Psychology, Vol.81, No.2, pp.275–280 (1969).",
 "[18] Wheeler, D. D.: Processes in word recognition, Cognitive Psychology, Vol.1, No.1, pp.59–85 (1970).",
]

ensure_cols(2)
for h, ps in SECS:
    lv = 2 if (len(h) > 2 and h[1] == '.' and h[2].isdigit()) else 1
    heading(h, lv)
    for pt in ps:
        para(pt)
heading("参考文献", 1)
for ref in REFS:
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.left_indent = Pt(18); pf.first_line_indent = Pt(-18); pf.space_after = Pt(2); pf.line_spacing = 1.1
    setrun(p.add_run(ref), size=8.5, jp=JP)

doc.save(DEST)
print("saved:", DEST)
print("sections:", len(SECS), "refs:", len(REFS), "doc sections:", len(doc.sections))
